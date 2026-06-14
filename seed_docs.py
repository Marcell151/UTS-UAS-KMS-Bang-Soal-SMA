import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header(doc):
    # Create a table for the header (1 row, 2 columns)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set column widths (Total 6.5 inches)
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.3)
    
    cell_logo = table.cell(0, 0)
    cell_text = table.cell(0, 1)
    
    # Add logo to left cell
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists('storage/logo.png'):
        r_logo = p_logo.add_run()
        r_logo.add_picture('storage/logo.png', width=Inches(0.9))
        
    # Add text to right cell
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p_text.add_run("SMA KRISTEN KALAM KUDUS MALANG\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102) # #003366
    
    run2 = p_text.add_run("NPSN: 20533621 – NSS: 304056102017 – NDS: E 32024002\n")
    run2.font.size = Pt(10)
    run3 = p_text.add_run("Prof. Moh. Yamin 47 Sukoharjo, Klojen, Malang 65118\n")
    run3.font.size = Pt(10)
    run4 = p_text.add_run("Telp 325829, 332980\n")
    run4.font.size = Pt(10)
    run5 = p_text.add_run("Website: https://www.skkkmalang.sch.id; E-mail: smakrkalamkudus@yahoo.com")
    run5.font.size = Pt(10)
    run5.font.color.rgb = RGBColor(0, 0, 255)
    
    # Add border paragraph outside table
    p_border = doc.add_paragraph()
    pBorder = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pBorder.append(bottom)
    p_border._p.get_or_add_pPr().append(pBorder)
    
    doc.add_paragraph()

files = [
    {
        "name": "Panduan_Penulisan_Soal_Pilihan_Ganda.docx",
        "title": "Panduan Penulisan Soal Pilihan Ganda (SOP)",
        "content": "Standar Operasional Prosedur (SOP) Penulisan Soal Pilihan Ganda\n\n1. TUJUAN\nMemberikan pedoman baku bagi tenaga pendidik di SMA Kristen Kalam Kudus Malang dalam menyusun butir soal pilihan ganda yang berkualitas, reliabel, dan valid.\n\n2. RUANG LINGKUP\nProsedur ini berlaku untuk penyusunan soal Penilaian Tengah Semester (PTS) dan Penilaian Akhir Semester (PAS).\n\n3. KETENTUAN UMUM\n- Pokok soal harus dirumuskan secara jelas dan tegas.\n- Pilihan jawaban harus homogen dan logis ditinjau dari segi materi.\n- Panjang rumusan pilihan jawaban harus relatif sama.\n- Pilihan jawaban tidak boleh mengandung pernyataan \"Semua pilihan jawaban di atas salah/benar\".\n- Soal harus menggunakan bahasa Indonesia yang baik dan benar sesuai EYD.\n\n4. PROSEDUR\n- Guru mata pelajaran menyusun kisi-kisi soal.\n- Penulisan butir soal berdasarkan indikator.\n- Validasi oleh Kepala Sekolah atau Koordinator Kurikulum.\n- Perbaikan soal (jika ada)."
    },
    {
        "name": "Template_Kisi_Kisi_Soal.docx",
        "title": "Template Kisi-Kisi Penulisan Soal",
        "content": "Format Standar Kisi-Kisi Penulisan Soal Ujian Sekolah\nSMA Kristen Kalam Kudus Malang\n\nMata Pelajaran : ...................\nKelas/Semester : ...................\nTahun Pelajaran: ...................\n\nInstruksi Penggunaan Template:\nSilakan isi tabel di bawah ini sesuai dengan silabus yang telah ditetapkan. Pastikan level kognitif sesuai dengan Taksonomi Bloom.\n\nTabel Kisi-Kisi:\nNo | Kompetensi Dasar | Materi | Indikator Soal | Level Kognitif (C1-C6) | Bentuk Soal | No. Soal"
    },
    {
        "name": "Soal_Matematika_Umum_X.docx",
        "title": "Bank Soal Matematika Umum Kelas X",
        "content": "Latihan Soal Eksponen dan Logaritma\n\nPilihlah satu jawaban yang paling tepat!\n\n1. Bentuk sederhana dari (a^3 b^-2 c) / (a b^-4 c^2) adalah...\nA. a^2 b^2 c^-1\nB. a^2 b^2 c\nC. a b c\nD. a^-2 b^-2 c\n\n2. Nilai dari 2log 8 + 3log 9 - 5log 25 adalah...\nA. 2\nB. 3\nC. 4\nD. 5"
    },
    {
        "name": "Soal_Fisika_XI_Dinamika.docx",
        "title": "Bank Soal Fisika Kelas XI: Dinamika Partikel",
        "content": "Evaluasi Bab Hukum Newton tentang Gerak\n\n1. Sebuah balok bermassa 10 kg berada di atas lantai kasar dengan koefisien gesekan statis 0.4 dan kinetis 0.2. Jika ditarik gaya 50 N mendatar, gaya gesek yang bekerja adalah... (g = 10 m/s^2)\nA. 20 N\nB. 40 N\nC. 50 N\nD. 100 N"
    },
    {
        "name": "Soal_Biologi_XII_Sel.docx",
        "title": "Bank Soal Biologi Kelas XII: Metabolisme",
        "content": "Penilaian Harian Metabolisme Sel\n\n1. Enzim merupakan biokatalisator yang mempercepat reaksi tanpa ikut bereaksi. Faktor yang TIDAK memengaruhi kerja enzim adalah...\nA. Suhu\nB. pH\nC. Konsentrasi enzim\nD. Warna substrat"
    },
    {
        "name": "Soal_Sejarah_XI_Kemerdekaan.docx",
        "title": "Bank Soal Sejarah Kemerdekaan",
        "content": "Soal Esai Sejarah Indonesia\n\n1. Analisis dampak Peristiwa Rengasdengklok terhadap Proklamasi Kemerdekaan RI.\n2. Sebutkan isi perjanjian Linggarjati."
    },
    {
        "name": "Soal_Bahasa_Inggris_X_Narrative.docx",
        "title": "English Question Bank: Narrative Text",
        "content": "Read the following text and answer the questions.\n\nOnce upon a time, there was a poor widow...\n\n1. What is the moral value of the story?\n2. Who is the main character?"
    },
    {
        "name": "Soal_Kimia_XI_Hidrokarbon.docx",
        "title": "Bank Soal Kimia Kelas XI: Hidrokarbon",
        "content": "Soal Latihan Hidrokarbon\n\n1. Senyawa alkana yang memiliki 5 atom karbon dinamakan...\nA. Metana\nB. Etana\nC. Propana\nD. Pentana"
    },
    {
        "name": "Soal_Sosiologi_XII_Perubahan.docx",
        "title": "Bank Soal Sosiologi Kelas XII: Perubahan Sosial",
        "content": "Soal Latihan Perubahan Sosial\n\n1. Salah satu faktor pendorong perubahan sosial dari luar (eksternal) adalah...\nA. Penemuan baru\nB. Konflik dalam masyarakat\nC. Pengaruh kebudayaan masyarakat lain\nD. Pemberontakan"
    },
    {
        "name": "Soal_Ekonomi_XI_Pajak.docx",
        "title": "Bank Soal Ekonomi Kelas XI: Perpajakan",
        "content": "Soal Latihan Pajak\n\n1. Pajak yang dipungut oleh pemerintah daerah disebut...\nA. Pajak Pusat\nB. Pajak Daerah\nC. Pajak Langsung\nD. Pajak Tidak Langsung"
    }
]

os.makedirs('upload', exist_ok=True)

for f in files:
    doc = Document()
    add_header(doc)
    
    h = doc.add_heading(f['title'], level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph(f['content'])
    
    doc.save(os.path.join('upload', f['name']))
    print(f"Generated upload/{f['name']}")

print("All documents generated successfully.")
